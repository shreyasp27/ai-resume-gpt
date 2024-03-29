import json
import google.generativeai as genai
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from io import BytesIO
from reportlab.lib.styles import ParagraphStyle
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from aws_lambda_powertools.utilities.data_classes import APIGatewayProxyEvent
import boto3
import os
from datetime import datetime

s3_client = boto3.client('s3')


# Initialize your genai model
model = genai.GenerativeModel(
    model_name="gemini-pro"
)

# Configure API key
google_api_key = os.environ.get('google_api_key')
genai.configure(api_key = google_api_key)

resume_prompt = (
    "Act as a professional resume writer. Following are your instructions:"
    "- Focus on updating/paraphrasing the Experience, Skills, and Project sections of the resume."
    "- Analyze the job description to identify key terms and skills."
    "- Incorporate these keywords into the Experience and Projects sections, ensuring they fit naturally and are relevant."
    "- Maintain the existing format of the resume. Do not alter the structure or style."
    "- Utilize only the information from the provided resume and 'about me' section for any updates."
    "- Do not add new content or information that isn’t already in the original resume or 'about me' section."
    "- Ensure the updated resume matches the job description keywords, enhancing its relevance for the role."
    "- Give the complete resume in the same format with the modified Experience, Skills, and Project sections of the resume."
    "- Use - for bullet points instead of *."
    "- Do not use bullet points for each skill."
    "- Make sure thatthere is proper spacing in the position name - city and company name - date is in the line below the position."
    "Resume:\n{resume}\n\n"
    "Job Description:\n{job_description}\n\n"
    "About Me:\n{about_me}\n"
)

cover_letter_prompt = (
    "Your task is to write a compelling cover letter for the job. Follow these instructions:\n\n"
    "- Use the information provided in the resume, job description, and about me sections to craft the letter.\n"
    "- The cover letter should be tailored to the job description, highlighting relevant skills and experiences.\n"
    "- Make sure the tone of the letter is professional and engaging.\n"
    "- The cover letter should complement the resume, not repeat its contents verbatim.\n"
    "- Focus on why the candidate is a good fit for the role and how they can contribute to the organization.\n"
    "- Make sure to address the cover letter from my name and address wherever necessary.\n\n"
    "- Let there be only one line gap between paragraphs.Do not have more than 3 paragraphs.\n\n"
    "- The letter should not be overly lengthy; keep it brief (3/4th page) but effective.\n\n"
    "- The letter should be concise, clear, and well-organized.\n\n"
    "Resume:\n{resume}\n\n"
    "Job Description:\n{job_description}\n\n"
    "About Me:\n{about_me}\n"
)

email_prompt = (
    "Compose a personalized email to the HR. Here are your instructions:\n\n"
    "- The email should introduce the candidate and express interest in the role described in the job description.\n"
    "- Use information from the provided resume and about me section to craft a personalized and relevant introduction.\n"
    "- Ensure the email is concise and to the point, while remaining engaging and professional.\n"
    "- Highlight key skills or experiences from the resume that align well with the job requirements.\n"
    "- Mention how the candidate's background and interests align with the company's values and mission, as described in the job description.\n"
    "- The email should not be overly lengthy; keep it brief but effective.\n\n"
    "Resume:\n{resume}\n\n"
    "Job Description:\n{job_description}\n\n"
    "About Me:\n{about_me}\n"
)


def update_resume(resume, job_description, about_me):
    formatted_prompt = resume_prompt.format(resume=resume, job_description=job_description, about_me=about_me)
    response = model.generate_content(formatted_prompt)
    if response.candidates:
        
        return response.candidates[0].content.parts[0].text

    return ""


def write_cover_letter(resume, job_description, about_me):
    formatted_prompt = cover_letter_prompt.format(resume=resume, job_description=job_description, about_me=about_me)
    response = model.generate_content(formatted_prompt)
    if response.candidates:
        # Get the first candidate's text directly
        return response.candidates[0].content.parts[0].text

    return ""

def write_personalized_email(resume, job_description, about_me):
    formatted_prompt = email_prompt.format(resume=resume, job_description=job_description, about_me=about_me)
    response = model.generate_content(formatted_prompt)
    if response.candidates:
        # Get the first candidate's text directly
        return response.candidates[0].content.parts[0].text

    return ""

def generate_pdf_buffer(text):
    buffer = BytesIO()
    styles = getSampleStyleSheet()
    
    
    bold_style = ParagraphStyle(
        'Bold',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        spaceAfter=6,
        leading=20,  
    )
    
    
    body_style = styles['BodyText']
    body_style.leading = 18  
    
    elements = []
    
    paragraphs = text.split("\n")
    for para in paragraphs:
        
        if '**' in para:
           
            para = para.replace('**', '')
            elements.append(Paragraph(para, bold_style))
        else:
            elements.append(Paragraph(para, body_style))
        elements.append(Spacer(1, 12))  
    
    # Build the PDF
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    doc.build(elements)
    
    buffer.seek(0)
    return buffer

def generate_docx_buffer(text):
    buffer = BytesIO()
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    for para in text.split("\n"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(para.replace('**', ''))
        run.bold = '**' in para
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def upload_to_s3(buffer, bucket_name, s3_filename):
    buffer.seek(0)
    s3_client.upload_fileobj(buffer, bucket_name, s3_filename)
    
    # Generate pre-signed URL for secure access
    url = s3_client.generate_presigned_url('get_object',
                                           Params={'Bucket': bucket_name, 'Key': s3_filename},
                                           ExpiresIn=3600)
    return url

def lambda_handler(event: APIGatewayProxyEvent, context):
    # Parse the event
    body = event.get('body', '{}')
     
    resume = body.get('resume', '')
    job_description = body.get('job_description', '')
    about_me = body.get('about_me', '')
   
    updated_resume_text = update_resume(resume, job_description, about_me)
    cover_letter_text = write_cover_letter(resume, job_description, about_me)
    personalized_email = write_personalized_email(resume, job_description, about_me)  
    
    pdf_buffer = generate_pdf_buffer(updated_resume_text)
    docx_buffer = generate_docx_buffer(cover_letter_text)
    

    bucket_name = 'ai-resume-gpt'
    unique_id = datetime.now().strftime("%Y%m%d%H%M%S")
    resume_key = f'resumes/resume_{unique_id}.pdf'
    cover_letter_key = f'cover_letters/cover_letter_{unique_id}.docx'
    
    # Upload the PDF resume and cover letter
    resume_url = upload_to_s3(pdf_buffer, bucket_name, resume_key)
    cover_letter_url = upload_to_s3(docx_buffer, bucket_name, cover_letter_key)
    
    response = {
        "statusCode": 200,
        "headers": {
            "Content-Type": "application/json"
        },
        "body": json.dumps({
            "resume_url": resume_url,
            "cover_letter_url": cover_letter_url,
            "personalized_email": personalized_email
        })
    }

    return response


